from __future__ import annotations

from copy import deepcopy
from pathlib import Path
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn

from .analyzer import analyze_docx
from .compiler import _clean_text, _split_contact, _split_entry_line
from .docx_parser import ParagraphModel, parse_docx


CANONICAL_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "corsair_canonical.docx"
SECTION_RE = re.compile(r"^[A-Z][A-Z &/]+$")


BLOCK_INDEXES = {
    "header_name": 0,
    "header_phone": 1,
    "header_contact": 2,
    "header_spacer": 3,
    "section_header": 4,
    "education_entry": 5,
    "education_detail": 6,
    "experience_entry": 13,
    "bullet": 14,
    "interests_line": 58,
}


def _body_element(document: Document):
    return document._body._element


def _clear_body(document: Document) -> None:
    body = _body_element(document)
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _append_cloned_paragraph(document: Document, template_paragraph):
    cloned = deepcopy(template_paragraph._p)
    body = _body_element(document)
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        body.insert(body.index(sect_pr), cloned)
    else:
        body.append(cloned)
    return document.paragraphs[-1]


def _clear_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:r"):
            paragraph._p.remove(child)
        if child.tag == qn("w:hyperlink"):
            paragraph._p.remove(child)


def _copy_run_format(source_run, target_run) -> None:
    if source_run._r.rPr is not None:
        target_run._r.insert(0, deepcopy(source_run._r.rPr))


def _add_formatted_run(paragraph, text: str, source_run, bold: Optional[bool] = None, italic: Optional[bool] = None):
    run = paragraph.add_run(text)
    _copy_run_format(source_run, run)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


class TemplateBlocks:
    def __init__(self, template_document: Document):
        self.document = template_document
        self.paragraphs = template_document.paragraphs

    def paragraph(self, name: str):
        return self.paragraphs[BLOCK_INDEXES[name]]

    def first_text_run(self, name: str):
        paragraph = self.paragraph(name)
        for run in paragraph.runs:
            if run.text:
                return run
        raise ValueError(f"No text run found for template block: {name}")

    def add_block(self, output_document: Document, name: str):
        paragraph = _append_cloned_paragraph(output_document, self.paragraph(name))
        _clear_runs(paragraph)
        return paragraph


def _is_section(paragraph: ParagraphModel) -> bool:
    return bool(SECTION_RE.fullmatch(_clean_text(paragraph.text)))


def _is_bullet(paragraph: ParagraphModel) -> bool:
    style = (paragraph.style_name or "").lower()
    return paragraph.numbering or "list" in style or _clean_text(paragraph.text).startswith(("• ", "● "))


def _is_source_only_subheading(text: str, current_section: Optional[str]) -> bool:
    if current_section in {None, "EDUCATION", "INTERESTS"}:
        return False
    cleaned = _clean_text(text)
    return cleaned.endswith(":") and "\t" not in cleaned


def _run_sources(blocks: TemplateBlocks, name: str):
    runs = [run for run in blocks.paragraph(name).runs if run.text]
    if not runs:
        raise ValueError(f"No runs found for block: {name}")
    return runs


def _write_simple_block(blocks: TemplateBlocks, document: Document, block_name: str, text: str) -> None:
    paragraph = blocks.add_block(document, block_name)
    _add_formatted_run(paragraph, text, blocks.first_text_run(block_name))


def _write_header(blocks: TemplateBlocks, document: Document, header_paragraphs: List[ParagraphModel]) -> None:
    _write_simple_block(blocks, document, "header_name", f"\t{_clean_text(header_paragraphs[0].text)}")

    contact_parts: List[str] = []
    for paragraph in header_paragraphs[1:]:
        contact_parts.extend(_split_contact(paragraph.text))

    if contact_parts:
        phone = contact_parts[1] if len(contact_parts) > 1 else contact_parts[0]
        _write_simple_block(blocks, document, "header_phone", f"\t{phone}")

    if len(contact_parts) >= 6:
        left, center, right = contact_parts[3], contact_parts[4], contact_parts[5]
    elif len(contact_parts) >= 5:
        left, center, right = contact_parts[0], contact_parts[3], contact_parts[4]
    elif len(contact_parts) >= 3:
        left, center, right = contact_parts[0], contact_parts[1], contact_parts[2]
    else:
        left, center, right = "", "", ""

    paragraph = blocks.add_block(document, "header_contact")
    run = blocks.first_text_run("header_contact")
    _add_formatted_run(paragraph, left, run)
    paragraph.add_run("\t")
    _add_formatted_run(paragraph, center, run)
    paragraph.add_run("\t")
    _add_formatted_run(paragraph, right, run)
    blocks.add_block(document, "header_spacer")


def _write_section(blocks: TemplateBlocks, document: Document, text: str) -> None:
    _write_simple_block(blocks, document, "section_header", _clean_text(text).upper())


def _write_entry(blocks: TemplateBlocks, document: Document, block_name: str, left: str, right: Optional[str]) -> None:
    paragraph = blocks.add_block(document, block_name)
    runs = _run_sources(blocks, block_name)
    normal_run = runs[1] if len(runs) > 1 else runs[0]
    italic_run = next((run for run in runs if run.italic), normal_run)
    bold_run = next((run for run in runs if run.bold), runs[0])

    pieces = [piece.strip() for piece in left.split(",")]
    if len(pieces) >= 3:
        _add_formatted_run(paragraph, pieces[0], bold_run, bold=True, italic=False)
        _add_formatted_run(paragraph, ", ", normal_run, bold=False, italic=False)
        _add_formatted_run(paragraph, pieces[1], italic_run, bold=False, italic=True)
        _add_formatted_run(paragraph, f", {', '.join(pieces[2:])}", normal_run, bold=False, italic=False)
    else:
        _add_formatted_run(paragraph, left, bold_run, bold=left.endswith(":"), italic=False)

    if right:
        paragraph.add_run("\t")
        _add_formatted_run(paragraph, right, normal_run, bold=False, italic=False)


def _write_labeled_detail(blocks: TemplateBlocks, document: Document, text: str) -> None:
    paragraph = blocks.add_block(document, "education_detail")
    runs = _run_sources(blocks, "education_detail")
    bold_run = next((run for run in runs if run.bold), runs[0])
    normal_run = next((run for run in runs if not run.bold), runs[-1])
    cleaned = _clean_text(text).removeprefix("• ").removeprefix("● ").strip()
    match = re.match(r"^(?P<label>[A-Za-z][A-Za-z /&]+)(?P<sep>:\s*)(?P<body>.+)$", cleaned)
    if not match:
        _add_formatted_run(paragraph, cleaned, normal_run, bold=False, italic=False)
        return
    _add_formatted_run(paragraph, match.group("label"), bold_run, bold=True, italic=False)
    _add_formatted_run(paragraph, match.group("sep"), normal_run, bold=False, italic=False)
    _add_formatted_run(paragraph, match.group("body"), normal_run, bold=False, italic=False)


def _write_bullet(blocks: TemplateBlocks, document: Document, text: str) -> None:
    paragraph = blocks.add_block(document, "bullet")
    run = blocks.first_text_run("bullet")
    _add_formatted_run(paragraph, _clean_text(text).removeprefix("• ").removeprefix("● ").strip(), run, bold=False, italic=False)


def compile_docx_from_blocks(input_path: str | Path, output_path: str | Path, render: bool = False) -> Dict[str, Any]:
    source = parse_docx(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document(CANONICAL_TEMPLATE_PATH)
    blocks = TemplateBlocks(Document(CANONICAL_TEMPLATE_PATH))
    _clear_body(document)

    paragraphs = [paragraph for paragraph in source.paragraphs if _clean_text(paragraph.text)]
    first_section_index = next((i for i, paragraph in enumerate(paragraphs) if _is_section(paragraph)), len(paragraphs))
    header_paragraphs = paragraphs[:first_section_index]
    body_paragraphs = paragraphs[first_section_index:]

    if header_paragraphs:
        _write_header(blocks, document, header_paragraphs)

    current_section: Optional[str] = None
    for paragraph in body_paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        if _is_section(paragraph):
            current_section = text.upper()
            _write_section(blocks, document, text)
            continue
        if _is_bullet(paragraph) and current_section == "EDUCATION":
            _write_labeled_detail(blocks, document, text)
            continue
        if _is_source_only_subheading(text, current_section):
            _write_bullet(blocks, document, text)
            continue
        if _is_bullet(paragraph):
            _write_bullet(blocks, document, text)
            continue
        if current_section == "INTERESTS":
            _write_simple_block(blocks, document, "interests_line", text)
            continue
        left, right = _split_entry_line(text)
        _write_entry(blocks, document, "education_entry" if current_section == "EDUCATION" else "experience_entry", left, right)

    document.save(output_path)

    before = analyze_docx(input_path, render=render)
    after = analyze_docx(output_path, render=render)
    return {
        "input_path": str(input_path),
        "output_path": str(output_path),
        "before": {
            "score": before["score"],
            "total_penalty": before["total_penalty"],
            "violation_count": len(before["violations"]),
        },
        "after": {
            "score": after["score"],
            "total_penalty": after["total_penalty"],
            "violation_count": len(after["violations"]),
        },
    }
