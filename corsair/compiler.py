from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from .analyzer import analyze_docx
from .docx_parser import ParagraphModel, parse_docx


SECTION_RE = re.compile(r"^[A-Z][A-Z &/]+$")
DATE_RE = re.compile(
    r"(?P<date>(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)"
    r"\s+\d{4}(?:\s*[-–]\s*(?:Present|Current|"
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{4}))?)$"
)
CANONICAL_FONT = "Garamond"
CANONICAL_TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "corsair_canonical.docx"


def _new_canonical_document() -> Document:
    document = Document(CANONICAL_TEMPLATE_PATH) if CANONICAL_TEMPLATE_PATH.exists() else Document()
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    return document


def _set_cell_border(paragraph, top: bool = False) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    if top:
        top_border = OxmlElement("w:top")
        top_border.set(qn("w:val"), "single")
        top_border.set(qn("w:sz"), "4")
        top_border.set(qn("w:space"), "2")
        top_border.set(qn("w:color"), "000000")
        p_bdr.append(top_border)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)


def _format_run(run, size: float = 10, bold: bool = False, italic: bool = False) -> None:
    run.font.name = CANONICAL_FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), CANONICAL_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), CANONICAL_FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = CANONICAL_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), CANONICAL_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CANONICAL_FONT)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1


def _clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"\bdd(?=[A-Z][a-z]+,\s*[A-Z]{2}\b)", "", text)
    text = re.sub(r" *\t *", "\t", text)
    text = re.sub(r"[ ]{2,}", " ", text)
    text = re.sub(r"\t{2,}", "\t", text)
    return text.strip()


def _split_contact(text: str) -> List[str]:
    text = _clean_text(text)
    rough_parts = [part.strip() for part in re.split(r"\t+| {2,}", text) if part.strip()]
    parts: List[str] = []
    email_re = re.compile(r"([\w.+-]+@[\w.-]+\.[A-Za-z]{2,})")
    for part in rough_parts:
        match = email_re.search(part)
        if not match:
            parts.append(part)
            continue
        before = part[: match.start()].strip(" |")
        email = match.group(1)
        after = part[match.end() :].strip(" |")
        if before:
            parts.append(before)
        parts.append(email)
        if after:
            parts.append(after)
    return parts


def _split_entry_line(text: str) -> Tuple[str, Optional[str]]:
    cleaned = _clean_text(text)
    tab_parts = [part.strip() for part in cleaned.split("\t") if part.strip()]
    if len(tab_parts) >= 2:
        return " ".join(tab_parts[:-1]).strip(), tab_parts[-1].strip()

    match = DATE_RE.search(cleaned)
    if not match:
        return cleaned, None
    left = cleaned[: match.start("date")].strip(" ,")
    return left, match.group("date").strip()


def _add_tight_paragraph(document: Document):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    return paragraph


def _add_header_name(document: Document, text: str) -> None:
    paragraph = _add_tight_paragraph(document)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(_clean_text(text))
    _format_run(run, size=16, bold=True)


def _add_center_contact_line(document: Document, text: str) -> None:
    paragraph = _add_tight_paragraph(document)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _format_run(run, size=10)


def _add_three_part_contact_line(document: Document, left: str, center: str, right: str) -> None:
    paragraph = _add_tight_paragraph(document)
    paragraph.paragraph_format.space_before = Pt(2.4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 0.75
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(3.75), WD_TAB_ALIGNMENT.CENTER)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    _format_run(paragraph.add_run(left), size=10)
    paragraph.add_run("\t")
    _format_run(paragraph.add_run(center), size=10)
    paragraph.add_run("\t")
    _format_run(paragraph.add_run(right), size=10)


def _add_header_spacer(document: Document) -> None:
    # Word renders the reference spacer as nearly zero-height because the
    # contact paragraph's bottom border carries the visual separation.
    return


def _add_header(document: Document, header_paragraphs: List[ParagraphModel]) -> None:
    _add_header_name(document, header_paragraphs[0].text)
    contact_parts: List[str] = []
    for paragraph in header_paragraphs[1:]:
        contact_parts.extend(_split_contact(paragraph.text))

    if len(contact_parts) >= 6:
        _add_center_contact_line(document, contact_parts[1])
        _add_three_part_contact_line(document, contact_parts[3], contact_parts[4], contact_parts[5])
        _add_header_spacer(document)
        return

    if len(contact_parts) >= 5:
        _add_center_contact_line(document, contact_parts[1])
        _add_three_part_contact_line(document, contact_parts[0], contact_parts[3], contact_parts[4])
        _add_header_spacer(document)
        return

    if len(contact_parts) == 4:
        _add_center_contact_line(document, contact_parts[1])
        _add_three_part_contact_line(document, contact_parts[0], contact_parts[2], contact_parts[3])
        _add_header_spacer(document)
        return

    for part in contact_parts:
        _add_center_contact_line(document, part)


def _add_section(document: Document, text: str, top_border: bool = False, space_before_pt: float = 0) -> None:
    paragraph = _add_tight_paragraph(document)
    paragraph.paragraph_format.space_before = Pt(space_before_pt)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 0.85
    _set_cell_border(paragraph, top=top_border)
    run = paragraph.add_run(_clean_text(text).upper())
    _format_run(run, size=10, bold=True)


def _add_styled_left_runs(paragraph, left: str) -> None:
    pieces = [piece.strip() for piece in left.split(",")]
    if len(pieces) >= 3:
        company = pieces[0]
        role = pieces[1]
        rest = ", ".join(pieces[2:])
        _format_run(paragraph.add_run(company), size=10, bold=True)
        _format_run(paragraph.add_run(", "), size=10)
        _format_run(paragraph.add_run(role), size=10, italic=True)
        _format_run(paragraph.add_run(f", {rest}"), size=10)
        return

    _format_run(paragraph.add_run(left), size=10, bold=left.endswith(":"))


def _add_labeled_text_runs(paragraph, text: str) -> None:
    match = re.match(r"^(?P<label>[A-Za-z][A-Za-z /&]+)(?P<sep>:\s*)(?P<body>.+)$", text)
    if not match:
        _format_run(paragraph.add_run(text), size=10)
        return
    _format_run(paragraph.add_run(match.group("label")), size=10, bold=True)
    _format_run(paragraph.add_run(match.group("sep")), size=10)
    _format_run(paragraph.add_run(match.group("body")), size=10)


def _add_entry_line(document: Document, left: str, right: Optional[str]) -> None:
    paragraph = _add_tight_paragraph(document)
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    _add_styled_left_runs(paragraph, left)
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right)
        _format_run(right_run, size=10)


def _add_plain_line(document: Document, text: str) -> None:
    paragraph = _add_tight_paragraph(document)
    run = paragraph.add_run(_clean_text(text))
    _format_run(run, size=10)


def _add_labeled_line(document: Document, text: str) -> None:
    paragraph = _add_tight_paragraph(document)
    _add_labeled_text_runs(paragraph, _clean_text(text).removeprefix("• ").removeprefix("● ").strip())


def _add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.48)
    paragraph.paragraph_format.first_line_indent = Inches(-0.23)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    bullet_run = paragraph.add_run("● ")
    _format_run(bullet_run, size=10)
    _add_labeled_text_runs(paragraph, _clean_text(text).removeprefix("• ").strip())


def _is_section(paragraph: ParagraphModel) -> bool:
    return bool(SECTION_RE.fullmatch(_clean_text(paragraph.text)))


def _is_bullet(paragraph: ParagraphModel) -> bool:
    style = (paragraph.style_name or "").lower()
    return paragraph.numbering or "list" in style or _clean_text(paragraph.text).startswith(("• ", "● "))


def compile_docx(input_path: str | Path, output_path: str | Path, render: bool = False) -> Dict[str, Any]:
    source = parse_docx(input_path)
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = _new_canonical_document()
    _configure_document(document)

    paragraphs = [paragraph for paragraph in source.paragraphs if _clean_text(paragraph.text)]
    first_section_index = next((i for i, paragraph in enumerate(paragraphs) if _is_section(paragraph)), len(paragraphs))
    header_paragraphs = paragraphs[:first_section_index]
    body_paragraphs = paragraphs[first_section_index:]

    if header_paragraphs:
        _add_header(document, header_paragraphs)

    current_section: Optional[str] = None
    for paragraph in body_paragraphs:
        text = _clean_text(paragraph.text)
        if not text:
            continue
        if _is_section(paragraph):
            current_section = text.upper()
            _add_section(
                document,
                text,
                top_border=current_section == "EDUCATION",
                space_before_pt=3.8 if current_section == "EDUCATION" else 0,
            )
            continue
        if _is_bullet(paragraph) and current_section == "EDUCATION":
            _add_labeled_line(document, text)
            continue
        if _is_bullet(paragraph):
            _add_bullet(document, text)
            continue
        if current_section == "INTERESTS":
            _add_plain_line(document, text)
            continue
        left, right = _split_entry_line(text)
        _add_entry_line(document, left, right)

    document.save(output_path)

    before = analyze_docx(input_path, render=render)
    after = analyze_docx(output_path, render=render)
    return {
        "input_path": str(input_path),
        "output_path": str(output_path),
        "before": {
            "score": before["score"],
            "total_penalty": before["total_penalty"],
            "violation_count": len(before["violations"]),
            "render_validation": before["document_summary"]["render_validation"],
        },
        "after": {
            "score": after["score"],
            "total_penalty": after["total_penalty"],
            "violation_count": len(after["violations"]),
            "render_validation": after["document_summary"]["render_validation"],
        },
    }
